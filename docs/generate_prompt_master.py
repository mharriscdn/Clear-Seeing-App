"""
Exports services/prompts/session.txt to docs/prompt_master.docx.

Run directly:  python docs/generate_prompt_master.py
Also called on every app startup (see app.py).
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SESSION_PROMPT_PATH = os.path.join(
    os.path.dirname(__file__), "..", "services", "prompts", "session.txt"
)
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "prompt_master.docx")


def generate():
    with open(SESSION_PROMPT_PATH, "r", encoding="utf-8") as f:
        content = f.read().strip()

    doc = Document()

    title = doc.add_heading("CLEAR SEEING GUIDE — Session Prompt v9", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    note = doc.add_paragraph(
        "Auto-generated. Do not edit manually. "
        "Run generate_prompt_master.py to update."
    )
    note.runs[0].italic = True
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    body = doc.add_paragraph()
    run = body.add_run(content)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.save(OUTPUT_PATH)
    print("prompt_master.docx written — v9 session prompt.")


if __name__ == "__main__":
    generate()
